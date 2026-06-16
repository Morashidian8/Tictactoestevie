#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
اپلیکیشن بازدید ایمنی و آتش‌نشانی ساختمان (HSE)
----------------------------------------------
- نمایش چک‌لیست به‌صورت ویزارد (یک ردیف در هر صفحه)
- ثبت وضعیت، توضیح و عکس برای هر ردیف
- خروجی گزارش تصویری Word (عکس‌ها + توضیح موارد دارای ایراد)
- خروجی Excel از ردیف‌های پرشده (با حذف موارد «کاربرد ندارد»)
"""

import io
import os
import json
import uuid
import zipfile
import datetime

from flask import (
    Flask, render_template, request, jsonify,
    send_file, abort,
)

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_DIR = os.path.join(BASE_DIR, "uploads")
DATA_DIR = os.path.join(BASE_DIR, "data")
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(DATA_DIR, exist_ok=True)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 64 * 1024 * 1024  # 64MB

# ------------------------------------------------------------------
# بارگذاری داده چک‌لیست
# ------------------------------------------------------------------
with open(os.path.join(BASE_DIR, "checklist_data.json"), encoding="utf-8") as f:
    CHECKLIST = json.load(f)

ITEMS = CHECKLIST["items"]
ROW_INDEX = {it["row"]: it for it in ITEMS}

# نگاشت وضعیت‌ها به نماد و امتیاز
STATUS_MAP = {
    "ok":   {"label": "مطلوب",          "symbol": "✅", "score": 3},
    "warn": {"label": "نیاز به بهبود",   "symbol": "⚠️", "score": 1},
    "bad":  {"label": "نامطلوب",         "symbol": "❌", "score": 0},
    "na":   {"label": "کاربرد ندارد",    "symbol": "N/A", "score": None},
}

ALLOWED_IMG = {".jpg", ".jpeg", ".png", ".gif", ".webp", ".bmp", ".heic"}


# ------------------------------------------------------------------
# صفحه اصلی (ویزارد)
# ------------------------------------------------------------------
@app.route("/")
def index():
    return render_template("index.html")


@app.route("/api/checklist")
def api_checklist():
    return jsonify(CHECKLIST)


# ------------------------------------------------------------------
# آپلود عکس برای یک ردیف
# ------------------------------------------------------------------
@app.route("/api/upload", methods=["POST"])
def api_upload():
    if "photo" not in request.files:
        return jsonify({"error": "no file"}), 400
    f = request.files["photo"]
    if not f.filename:
        return jsonify({"error": "empty filename"}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ALLOWED_IMG:
        return jsonify({"error": "نوع فایل مجاز نیست"}), 400
    name = f"{uuid.uuid4().hex}{ext}"
    f.save(os.path.join(UPLOAD_DIR, name))
    return jsonify({"filename": name})


@app.route("/uploads/<path:name>")
def uploaded_file(name):
    path = os.path.join(UPLOAD_DIR, name)
    if not os.path.isfile(path):
        abort(404)
    return send_file(path)


# ------------------------------------------------------------------
# ذخیره موقت کل پاسخ‌ها (برای ادامه کار بعدی)
# ------------------------------------------------------------------
@app.route("/api/save", methods=["POST"])
def api_save():
    payload = request.get_json(force=True)
    with open(os.path.join(DATA_DIR, "session.json"), "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)
    return jsonify({"ok": True})


@app.route("/api/load")
def api_load():
    path = os.path.join(DATA_DIR, "session.json")
    if not os.path.isfile(path):
        return jsonify({})
    with open(path, encoding="utf-8") as f:
        return jsonify(json.load(f))


# ------------------------------------------------------------------
# کمکی: پردازش payload ارسالی از فرانت
# ------------------------------------------------------------------
def _normalize(payload):
    building = payload.get("building", {})
    answers = payload.get("answers", {})
    rows = []
    for it in ITEMS:
        r = str(it["row"])
        ans = answers.get(r) or answers.get(it["row"]) or {}
        status = ans.get("status")
        if not status:
            continue  # پاسخ داده نشده
        sm = STATUS_MAP.get(status, {})
        rows.append({
            "row": it["row"],
            "section": it["section"],
            "item": it["item"],
            "ref": it.get("ref", ""),
            "priority": it.get("priority", ""),
            "status": status,
            "status_label": sm.get("label", ""),
            "status_symbol": sm.get("symbol", ""),
            "score": sm.get("score"),
            "note": (ans.get("note") or "").strip(),
            "photos": ans.get("photos", []) or [],
        })
    return building, rows


# ------------------------------------------------------------------
# خروجی Excel (حذف موارد «کاربرد ندارد»)
# ------------------------------------------------------------------
@app.route("/api/export/excel", methods=["POST"])
def export_excel():
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    building, rows = _normalize(request.get_json(force=True))

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "گزارش بازدید"
    ws.sheet_view.rightToLeft = True

    thin = Side(style="thin", color="B0B0B0")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)
    center = Alignment(horizontal="center", vertical="center", wrap_text=True)
    right = Alignment(horizontal="right", vertical="center", wrap_text=True)

    # هدر اطلاعات ساختمان
    ws.merge_cells("A1:H1")
    c = ws["A1"]
    c.value = "گزارش بازدید ایمنی و آتش‌نشانی ساختمان"
    c.font = Font(bold=True, size=14, color="FFFFFF")
    c.fill = PatternFill("solid", fgColor="1F4E78")
    c.alignment = center
    ws.row_dimensions[1].height = 28

    info = [
        ("نام ساختمان / پروژه", building.get("name", "")),
        ("آدرس", building.get("address", "")),
        ("تاریخ بازدید", building.get("date", "")),
        ("بازرس HSE", building.get("inspector", "")),
        ("کد ساختمان", building.get("code", "")),
        ("دوره بازدید", building.get("period", "")),
    ]
    r = 2
    for i in range(0, len(info), 2):
        ws.cell(r, 1, info[i][0]).font = Font(bold=True)
        ws.cell(r, 2, info[i][1])
        if i + 1 < len(info):
            ws.cell(r, 4, info[i + 1][0]).font = Font(bold=True)
            ws.cell(r, 5, info[i + 1][1])
        r += 1

    # سرستون جدول
    headers = ["ردیف", "حوزه", "شرح آیتم", "مرجع قانونی",
               "وضعیت", "مغایرت / توضیح", "اولویت", "امتیاز", "تعداد عکس"]
    head_row = r + 1
    for col, h in enumerate(headers, 1):
        cell = ws.cell(head_row, col, h)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="2E75B6")
        cell.alignment = center
        cell.border = border

    fill_map = {"ok": "C6EFCE", "warn": "FFEB9C", "bad": "FFC7CE"}

    rr = head_row + 1
    scored = []
    for row in rows:
        if row["status"] == "na":  # حذف موارد کاربرد ندارد از خروجی اکسل
            continue
        vals = [
            row["row"], row["section"], row["item"], row["ref"],
            f'{row["status_symbol"]} {row["status_label"]}',
            row["note"], row["priority"],
            row["score"] if row["score"] is not None else "",
            len(row["photos"]),
        ]
        for col, v in enumerate(vals, 1):
            cell = ws.cell(rr, col, v)
            cell.border = border
            cell.alignment = right if col in (2, 3, 4, 6) else center
        fc = fill_map.get(row["status"])
        if fc:
            ws.cell(rr, 5).fill = PatternFill("solid", fgColor=fc)
        if row["score"] is not None:
            scored.append(row["score"])
        rr += 1

    # ردیف جمع و درصد
    if scored:
        total = sum(scored)
        maxp = len(scored) * 3
        pct = round(total / maxp * 100, 1) if maxp else 0
        ws.cell(rr + 1, 3, "جمع امتیاز کسب‌شده").font = Font(bold=True)
        ws.cell(rr + 1, 8, total).font = Font(bold=True)
        ws.cell(rr + 2, 3, "حداکثر امتیاز ممکن").font = Font(bold=True)
        ws.cell(rr + 2, 8, maxp).font = Font(bold=True)
        ws.cell(rr + 3, 3, "درصد ایمنی").font = Font(bold=True)
        pc = ws.cell(rr + 3, 8, f"{pct}%")
        pc.font = Font(bold=True)
        level, color = _safety_level(pct)
        lv = ws.cell(rr + 3, 4, level)
        lv.font = Font(bold=True)
        lv.fill = PatternFill("solid", fgColor=color)

    widths = [7, 26, 45, 20, 16, 35, 10, 9, 10]
    for col, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = w

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    fname = f"HSE_Report_{datetime.date.today()}.xlsx"
    return send_file(bio, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def _safety_level(pct):
    if pct >= 85:
        return "ایمن – قابل قبول", "C6EFCE"
    if pct >= 70:
        return "متوسط – نیاز به توجه", "FFEB9C"
    if pct >= 50:
        return "ضعیف – اقدام فوری", "FCD5B4"
    return "بحرانی – مداخله اضطراری", "FFC7CE"


# ------------------------------------------------------------------
# خروجی گزارش تصویری Word
# ------------------------------------------------------------------
@app.route("/api/export/word", methods=["POST"])
def export_word():
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    building, rows = _normalize(request.get_json(force=True))
    doc = Document()

    # راست‌چین کردن کل سند
    def rtl(par):
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = par._p.get_or_add_pPr()
        bidi = OxmlElement("w:bidi")
        pPr.append(bidi)
        return par

    def set_font(run, size=11, bold=False, color=None):
        run.font.name = "B Nazanin"
        run.font.size = Pt(size)
        run.font.bold = bold
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:cs"), "B Nazanin")
        if color:
            run.font.color.rgb = color

    # عنوان
    h = rtl(doc.add_paragraph())
    set_font(h.add_run("گزارش تصویری بازدید ایمنی و آتش‌نشانی ساختمان"), 18, True,
             RGBColor(0x1F, 0x4E, 0x78))

    # اطلاعات ساختمان
    info = [
        ("نام ساختمان / پروژه", building.get("name", "")),
        ("آدرس", building.get("address", "")),
        ("تاریخ بازدید", building.get("date", "")),
        ("بازرس HSE", building.get("inspector", "")),
        ("کد ساختمان", building.get("code", "")),
        ("دوره بازدید", building.get("period", "")),
    ]
    for k, v in info:
        if not v:
            continue
        p = rtl(doc.add_paragraph())
        set_font(p.add_run(f"{k}: "), 12, True)
        set_font(p.add_run(str(v)), 12, False)

    # فقط موارد دارای ایراد (نامطلوب/نیاز به بهبود) یا دارای عکس/توضیح
    issues = [r for r in rows
              if r["status"] in ("bad", "warn") or r["photos"] or r["note"]]
    issues = [r for r in issues if r["status"] != "na"]

    p = rtl(doc.add_paragraph())
    set_font(p.add_run(f"تعداد موارد دارای ایراد / نیازمند توجه: {len(issues)}"),
             13, True, RGBColor(0xC0, 0x00, 0x00))

    doc.add_paragraph()

    for idx, r in enumerate(issues, 1):
        # عنوان مورد
        p = rtl(doc.add_paragraph())
        set_font(p.add_run(f"{idx}) ردیف {r['row']} – {r['section']}"), 13, True,
                 RGBColor(0x1F, 0x4E, 0x78))

        p = rtl(doc.add_paragraph())
        set_font(p.add_run("شرح آیتم: "), 12, True)
        set_font(p.add_run(r["item"]), 12)

        color = {"bad": RGBColor(0xC0, 0x00, 0x00),
                 "warn": RGBColor(0xBF, 0x8F, 0x00),
                 "ok": RGBColor(0x00, 0x80, 0x00)}.get(r["status"], None)
        p = rtl(doc.add_paragraph())
        set_font(p.add_run("وضعیت: "), 12, True)
        set_font(p.add_run(f'{r["status_symbol"]} {r["status_label"]}'), 12, True, color)
        if r["priority"]:
            set_font(p.add_run(f"   |   اولویت اصلاح: {r['priority']}"), 12, False)

        if r["note"]:
            p = rtl(doc.add_paragraph())
            set_font(p.add_run("توضیح ناظر: "), 12, True)
            set_font(p.add_run(r["note"]), 12)

        # عکس‌ها
        for ph in r["photos"]:
            path = os.path.join(UPLOAD_DIR, ph)
            if os.path.isfile(path):
                pic = doc.add_paragraph()
                pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    pic.add_run().add_picture(path, width=Inches(4.2))
                except Exception:
                    pass
                cap = rtl(doc.add_paragraph())
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = r["note"] if r["note"] else r["item"]
                set_font(cap.add_run(f"تصویر مورد {r['row']}: {caption}"), 10, False,
                         RGBColor(0x60, 0x60, 0x60))

        # خط جداکننده
        sep = doc.add_paragraph()
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(sep.add_run("ـ" * 30), 10, False, RGBColor(0xC0, 0xC0, 0xC0))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    fname = f"HSE_Visual_Report_{datetime.date.today()}.docx"
    return send_file(bio, as_attachment=True, download_name=fname,
                     mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=True)
